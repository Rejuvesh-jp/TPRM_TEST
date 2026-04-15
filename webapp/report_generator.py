"""
Report Generator
=================
Generate PDF and Word reports from assessment data.
"""
import io
import logging
from datetime import datetime
from pathlib import Path

logger = logging.getLogger("tprm.reports")

# Logo path — shared between PDF and Word generators
LOGO_PATH = Path(__file__).resolve().parent.parent / "inputs" / "image.png"


# ────────────────────────────────────────────────────────────
#  WORD REPORT
# ────────────────────────────────────────────────────────────

def generate_word(report: dict, vendor_name: str) -> io.BytesIO:
    """Generate an enterprise-grade Word document report."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()
    now_str = datetime.now().strftime("%B %d, %Y  %I:%M %p")

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── Style definitions ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    INDIGO = RGBColor(0x4F, 0x46, 0xE5)
    DARK = RGBColor(0x1E, 0x29, 0x3B)
    GRAY = RGBColor(0x64, 0x74, 0x8B)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    severity_colors = {
        "critical": RGBColor(0xDC, 0x26, 0x26),
        "high": RGBColor(0xEA, 0x58, 0x0C),
        "medium": RGBColor(0xD9, 0x77, 0x06),
        "low": RGBColor(0x05, 0x96, 0x69),
    }

    def _set_cell_shading(cell, hex_color):
        """Apply background shading to a table cell."""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _style_header_row(row, col_count):
        """Apply indigo background + white bold text to a header row."""
        for i in range(col_count):
            cell = row.cells[i]
            _set_cell_shading(cell, "4F46E5")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    def _clean_xml(text):
        """Remove characters illegal in XML (NULL bytes and control chars except tab/newline/CR)."""
        import re
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', str(text))

    def _add_body_cell(cell, text, font_size=Pt(9), bold=False, color=None, align=None):
        """Write text into a table cell with formatting."""
        cell.text = ""
        p = cell.paragraphs[0]
        if align:
            p.alignment = align
        run = p.add_run(_clean_xml(text))
        run.font.size = font_size
        run.font.name = "Calibri"
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    def _set_table_borders(table):
        """Set enterprise-grade borders: thick navy outer frame, solid inner lines."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top    w:val="single" w:sz="12" w:space="0" w:color="1E293B"/>'
            '  <w:left   w:val="single" w:sz="12" w:space="0" w:color="1E293B"/>'
            '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="1E293B"/>'
            '  <w:right  w:val="single" w:sz="12" w:space="0" w:color="1E293B"/>'
            '  <w:insideH w:val="single" w:sz="6"  w:space="0" w:color="CBD5E1"/>'
            '  <w:insideV w:val="single" w:sz="6"  w:space="0" w:color="CBD5E1"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

    def _add_section_heading(text):
        """Add a styled section heading."""
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.color.rgb = INDIGO
        run.font.bold = True
        # Add a thin indigo line below
        border_p = doc.add_paragraph()
        border_p.space_after = Pt(8)
        pPr = border_p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="4F46E5"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)

    # ═══════════════════════════════════════════════════════
    #  COVER PAGE
    # ═══════════════════════════════════════════════════════

    # Logo
    if LOGO_PATH.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.space_after = Pt(20)
        logo_p.add_run().add_picture(str(LOGO_PATH), width=Inches(2.0))

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.space_before = Pt(30)
    title_p.space_after = Pt(6)
    run = title_p.add_run("TPRM AI Assessment Report")
    run.font.size = Pt(26)
    run.font.color.rgb = INDIGO
    run.font.bold = True

    # Vendor subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.space_after = Pt(4)
    run = sub_p.add_run(f"Vendor: {vendor_name}")
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    # Date line
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.space_after = Pt(30)
    run = date_p.add_run(f"Report Generated: {now_str}")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    # ── Summary info table (cover page) ──
    summary = report.get("summary", {})
    meta = report.get("meta", {})  # meta may be injected by caller if needed
    division = meta.get("division", "")
    nature_of_engagement = meta.get("nature_of_engagement", "")
    pre_assessment_scores = meta.get("pre_assessment_scores")
    noe_display = nature_of_engagement.title() if nature_of_engagement else ""
    if pre_assessment_scores:
        noe_display += f" (Score: {pre_assessment_scores.get('total_score', 0)})"
    risk_rating = meta.get("risk_rating") or report.get("risk_rating", {})
    if isinstance(risk_rating, dict):
        overall = (risk_rating.get("overall") or "").upper()
    else:
        overall = (risk_rating or "").upper()
    # If not assigned, show blank
    if not overall or overall == "N/A":
        overall = ""

    # Only open gaps
    all_gaps = report.get("gaps", [])
    open_gaps = [g for g in all_gaps if g.get("gap_status", "open") == "open"]

    info_rows = [
        ("Division", division),
        ("Nature of Engagement", noe_display),
        ("Vendor Risk Rating", overall),
        ("Open Gaps", str(len(open_gaps))),
        ("Total Remedial Actions", str(summary.get("total_remedial_actions", len(report.get("remedial_plan", []))))),
        ("Total Recommendations", str(summary.get("total_recommendations", 0))),
        ("Report Date", now_str),
    ]

    info_table = doc.add_table(rows=len(info_rows), cols=2, style="Table Grid")
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = True
    _set_table_borders(info_table)

    for idx, (label, value) in enumerate(info_rows):
        _set_cell_shading(info_table.rows[idx].cells[0], "4F46E5")
        _add_body_cell(info_table.rows[idx].cells[0], label, bold=True, color=WHITE,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_shading(info_table.rows[idx].cells[1], "F8FAFC")
        _add_body_cell(info_table.rows[idx].cells[1], value, bold=False,
                       align=WD_ALIGN_PARAGRAPH.LEFT)

    # Set column widths for info table
    for row in info_table.rows:
        row.cells[0].width = Cm(6.5)
        row.cells[1].width = Cm(10.5)

    # ── Executive Summary ──
    exec_summary = summary.get("executive_summary", "")
    if exec_summary:
        doc.add_paragraph()  # spacer
        _add_section_heading("Executive Summary")
        es_p = doc.add_paragraph()
        es_p.space_after = Pt(12)
        run = es_p.add_run(exec_summary)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    #  SECTION 1: GAPS (only open)
    # ═══════════════════════════════════════════════════════

    _add_section_heading(f"Identified Gaps ({len(open_gaps)})")

    if not open_gaps:
        p = doc.add_paragraph()
        run = p.add_run("No open gaps identified.")
        run.font.color.rgb = GRAY
    else:
        gap_headers = ["#", "Gap Type", "Description", "Evidence"]
        gap_table = doc.add_table(rows=1 + len(open_gaps), cols=4, style="Table Grid")
        gap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        gap_table.autofit = False
        _set_table_borders(gap_table)
        for ci, h in enumerate(gap_headers):
            gap_table.rows[0].cells[ci].text = h
        _style_header_row(gap_table.rows[0], 4)
        col_widths = [Cm(1), Cm(3.5), Cm(9), Cm(3.5)]
        for row in gap_table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = w
        for i, g in enumerate(open_gaps, 1):
            row = gap_table.rows[i]
            _add_body_cell(row.cells[0], str(i), align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_body_cell(row.cells[1], g.get("gap_type", ""))
            _add_body_cell(row.cells[2], g.get("description", ""))
            _add_body_cell(row.cells[3], g.get("evidence_assessment", "—"))
            if i % 2 == 0:
                for ci in range(4):
                    _set_cell_shading(row.cells[ci], "F8FAFC")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    #  SECTION 1.5: REMEDIAL ACTION PLAN
    # ═══════════════════════════════════════════════════════

    remedial_actions = report.get("remedial_plan", [])
    _add_section_heading(f"Remedial Action Plan ({len(remedial_actions)})")

    if remedial_actions:
        rem_headers = ["#", "Priority", "Action", "Timeline", "Owner", "Acceptance Criteria"]
        rem_table = doc.add_table(rows=1 + len(remedial_actions), cols=6, style="Table Grid")
        rem_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rem_table.autofit = False
        _set_table_borders(rem_table)

        for ci, h in enumerate(rem_headers):
            rem_table.rows[0].cells[ci].text = h
        _style_header_row(rem_table.rows[0], 6)

        # #=0.8cm, Priority=2.5cm, Action=6cm, Timeline=2cm, Owner=3cm, Criteria=3cm
        rem_col_widths = [Cm(0.8), Cm(2.5), Cm(6), Cm(2), Cm(3), Cm(3)]
        for row in rem_table.rows:
            for ci, w in enumerate(rem_col_widths):
                row.cells[ci].width = w

        priority_colors = {
            "immediate": RGBColor(0xDC, 0x26, 0x26),
            "short_term": RGBColor(0xEA, 0x58, 0x0C),
            "medium_term": RGBColor(0x0E, 0xA5, 0xE9),
            "long_term": RGBColor(0x64, 0x74, 0x8B),
        }
        priority_bg = {
            "immediate": "FEE2E2",
            "short_term": "FFEDD5",
            "medium_term": "E0F2FE",
            "long_term": "F1F5F9",
        }

        for i, act in enumerate(remedial_actions, 1):
            pri = (act.get("priority") or "medium_term").lower()
            pri_label = pri.replace("_", " ").upper()
            row = rem_table.rows[i]
            _add_body_cell(row.cells[0], str(i), align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_body_cell(row.cells[1], pri_label, bold=True,
                           color=priority_colors.get(pri, DARK),
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_body_cell(row.cells[2], act.get("action", ""))
            _add_body_cell(row.cells[3], act.get("timeline", "—"), align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_body_cell(row.cells[4], act.get("owner", "—"))
            _add_body_cell(row.cells[5], act.get("acceptance_criteria", "—"))

            bg = priority_bg.get(pri)
            if bg:
                _set_cell_shading(row.cells[1], bg)
            if i % 2 == 0:
                for ci in range(6):
                    if ci != 1:
                        _set_cell_shading(row.cells[ci], "F8FAFC")
    else:
        p = doc.add_paragraph()
        run = p.add_run("No remedial actions generated.")
        run.font.color.rgb = GRAY

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    #  SECTION 2: RECOMMENDED CONTRACT CLAUSES
    # ═══════════════════════════════════════════════════════

    recs = report.get("recommendations", [])
    _add_section_heading(f"Recommended Contract Clauses ({len(recs)})")

    if recs:
        rec_headers = ["#", "Type", "Priority", "Clause Text", "Justification", "Coverage"]
        rec_table = doc.add_table(rows=1 + len(recs), cols=6, style="Table Grid")
        rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rec_table.autofit = False
        _set_table_borders(rec_table)

        for ci, h in enumerate(rec_headers):
            rec_table.rows[0].cells[ci].text = h
        _style_header_row(rec_table.rows[0], 6)

        # #=0.8cm, Type=1.8cm, Priority=2cm, Clause=6.5cm, Justification=4cm, Coverage=1.9cm
        col_widths = [Cm(0.8), Cm(1.8), Cm(2), Cm(6.5), Cm(4), Cm(1.9)]
        for row in rec_table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = w

        for i, rec in enumerate(recs, 1):
            source = rec.get("source", "")
            if source == "existing":
                rec_type = "Existing"
            elif source == "new":
                rec_type = "New"
            else:
                rec_type = "Rec."
            priority = (rec.get("priority") or "").replace("_", " ").title()

            row = rec_table.rows[i]
            _add_body_cell(row.cells[0], str(i), align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_body_cell(row.cells[1], rec_type, align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_body_cell(row.cells[2], priority, align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_body_cell(row.cells[3], rec.get("clause_text", ""))
            _add_body_cell(row.cells[4], rec.get("justification", ""))
            _add_body_cell(row.cells[5], rec.get("existing_coverage", "N/A"),
                           align=WD_ALIGN_PARAGRAPH.CENTER)

            if i % 2 == 0:
                for ci in range(6):
                    _set_cell_shading(row.cells[ci], "F8FAFC")
    else:
        p = doc.add_paragraph()
        run = p.add_run("No recommendations generated.")
        run.font.color.rgb = GRAY

    # ── Footer note ──
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"— End of Report — Generated on {now_str} —")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # ── Write to buffer ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ────────────────────────────────────────────────────────────
#  PDF REPORT
# ────────────────────────────────────────────────────────────

def generate_pdf(report: dict, vendor_name: str) -> io.BytesIO:
    """Generate an enterprise-grade PDF report with logo, timestamps, gaps, and recommended clauses."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, Image,
    )

    now_str = datetime.now().strftime("%B %d, %Y  %I:%M %p")
    buf = io.BytesIO()

    def _footer_header(canvas, doc_template):
        """Draw page header (logo) and footer (date + page number) on every page."""
        canvas.saveState()
        # Footer
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(HexColor("#64748B"))
        canvas.drawString(20 * mm, 10 * mm,
                          f"TPRM AI Assessment Report  |  {vendor_name}  |  {now_str}")
        canvas.drawRightString(A4[0] - 20 * mm, 10 * mm,
                               f"Page {canvas.getPageNumber()}")
        # Thin header line
        canvas.setStrokeColor(HexColor("#4F46E5"))
        canvas.setLineWidth(0.5)
        canvas.line(20 * mm, A4[1] - 22 * mm, A4[0] - 20 * mm, A4[1] - 22 * mm)
        # Thin footer line
        canvas.line(20 * mm, 14 * mm, A4[0] - 20 * mm, 14 * mm)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=25 * mm, bottomMargin=20 * mm,
    )

    # ── Colors ──
    INDIGO = HexColor("#4F46E5")
    DARK = HexColor("#1E293B")
    GRAY = HexColor("#64748B")
    WHITE = HexColor("#FFFFFF")
    LIGHT_BG = HexColor("#F8FAFC")
    CRITICAL_BG = HexColor("#FEE2E2")
    HIGH_BG = HexColor("#FFEDD5")
    MEDIUM_BG = HexColor("#FEF9C3")
    LOW_BG = HexColor("#D1FAE5")

    severity_colors = {
        "critical": (HexColor("#DC2626"), CRITICAL_BG),
        "high": (HexColor("#EA580C"), HIGH_BG),
        "medium": (HexColor("#D97706"), MEDIUM_BG),
        "low": (HexColor("#059669"), LOW_BG),
    }

    # ── Styles ──
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("Title2", parent=styles["Title"], fontSize=24,
                              textColor=INDIGO, spaceAfter=4, alignment=TA_CENTER))
    styles.add(ParagraphStyle("Subtitle", parent=styles["Normal"], fontSize=12,
                              textColor=GRAY, spaceAfter=4, alignment=TA_CENTER))
    styles.add(ParagraphStyle("DateLine", parent=styles["Normal"], fontSize=9,
                              textColor=GRAY, spaceAfter=14, alignment=TA_CENTER))
    styles.add(ParagraphStyle("SectionHead", parent=styles["Heading2"], fontSize=14,
                              textColor=INDIGO, spaceBefore=16, spaceAfter=8))
    styles.add(ParagraphStyle("CellText", parent=styles["Normal"], fontSize=8,
                              leading=11, textColor=DARK))
    styles.add(ParagraphStyle("CellBold", parent=styles["Normal"], fontSize=8,
                              leading=11, textColor=WHITE, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle("Body", parent=styles["Normal"], fontSize=10,
                              leading=14, textColor=DARK, spaceAfter=6))
    styles.add(ParagraphStyle("SmallGray", parent=styles["Normal"], fontSize=8,
                              textColor=GRAY, alignment=TA_CENTER))

    elements = []
    page_width = A4[0] - 40 * mm  # usable width

    # ── Title Page ──
    elements.append(Spacer(1, 15 * mm))

    # Logo
    if LOGO_PATH.exists():
        logo = Image(str(LOGO_PATH), width=60 * mm, height=60 * mm)
        logo.hAlign = "CENTER"
        elements.append(logo)
        elements.append(Spacer(1, 10 * mm))

    elements.append(Paragraph("TPRM AI Assessment Report", styles["Title2"]))
    elements.append(Paragraph(f"Vendor: {vendor_name}", styles["Subtitle"]))
    elements.append(Paragraph(f"Report Generated: {now_str}", styles["DateLine"]))

    summary = report.get("summary", {})
    meta = report.get("meta", {})
    division = meta.get("division", "")
    nature_of_engagement = meta.get("nature_of_engagement", "")
    pre_assessment_scores_pdf = meta.get("pre_assessment_scores")
    noe_display_pdf = nature_of_engagement.title() if nature_of_engagement else ""
    if pre_assessment_scores_pdf:
        noe_display_pdf += f" (Score: {pre_assessment_scores_pdf.get('total_score', 0)})"
    risk_rating = meta.get("risk_rating") or report.get("risk_rating", {})
    if isinstance(risk_rating, dict):
        overall = (risk_rating.get("overall") or "").upper()
    else:
        overall = (risk_rating or "").upper()
    # If not assigned, show blank
    if not overall or overall == "N/A":
        overall = ""

    # Only open gaps
    all_gaps = report.get("gaps", [])
    open_gaps = [g for g in all_gaps if g.get("gap_status", "open") == "open"]

    elements.append(Spacer(1, 8 * mm))
    # Summary info table
    info_data = [
        ["Division", division],
        ["Nature of Engagement", noe_display_pdf],
        ["Vendor Risk Rating", overall],
        ["Open Gaps", str(len(open_gaps))],
        ["Total Recommendations", str(summary.get("total_recommendations", 0))],
        ["Report Date", now_str],
    ]
    info_table = Table(info_data, colWidths=[page_width * 0.4, page_width * 0.6])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), INDIGO),
        ("TEXTCOLOR", (0, 0), (0, -1), WHITE),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#E2E8F0")),
        ("BACKGROUND", (1, 0), (1, -1), LIGHT_BG),
    ]))
    elements.append(info_table)

    # Executive summary
    exec_summary = summary.get("executive_summary", "")
    if exec_summary:
        elements.append(Spacer(1, 8 * mm))
        elements.append(Paragraph("Executive Summary", styles["SectionHead"]))
        elements.append(Paragraph(exec_summary, styles["Body"]))

    elements.append(PageBreak())

    # ── Helper: build data tables ──
    def _sev_cell(text, level):
        """Return colored severity paragraph."""
        color, _ = severity_colors.get(level.lower(), (DARK, LIGHT_BG))
        return Paragraph(f'<font color="{color.hexval()}">{text.upper()}</font>',
                         styles["CellText"])

    header_style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), INDIGO),
        ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("GRID", (0, 0), (-1, -1), 0.4, HexColor("#E2E8F0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT_BG]),
    ])

    # ── Section 1: Gaps (only open) ──
    elements.append(Paragraph(f"Identified Gaps ({len(open_gaps)})", styles["SectionHead"]))
    elements.append(HRFlowable(width="100%", thickness=1, color=INDIGO, spaceAfter=6))
    if open_gaps:
        col_widths = [page_width * w for w in [0.05, 0.17, 0.48, 0.30]]
        gap_data = [
            [Paragraph(h, styles["CellBold"]) for h in ["#", "Type", "Description", "Evidence"]]
        ]
        for i, g in enumerate(open_gaps, 1):
            gap_data.append([
                Paragraph(str(i), styles["CellText"]),
                Paragraph(g.get("gap_type", ""), styles["CellText"]),
                Paragraph(g.get("description", ""), styles["CellText"]),
                Paragraph(g.get("evidence_assessment", "—"), styles["CellText"]),
            ])
        t = Table(gap_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(header_style)
        elements.append(t)
    else:
        elements.append(Paragraph("No open gaps identified.", styles["Body"]))

    elements.append(PageBreak())

    # ── Section 2: Recommended Clauses ──
    recs = report.get("recommendations", [])
    elements.append(Paragraph(f"Recommended Contract Clauses ({len(recs)})", styles["SectionHead"]))
    elements.append(HRFlowable(width="100%", thickness=1, color=INDIGO, spaceAfter=6))

    if recs:
        col_widths = [page_width * w for w in [0.04, 0.10, 0.11, 0.42, 0.25, 0.08]]
        rec_data = [
            [Paragraph(h, styles["CellBold"]) for h in
             ["#", "Type", "Priority", "Clause Text", "Justification", "Coverage"]]
        ]
        for i, rec in enumerate(recs, 1):
            source = rec.get("source", "")
            if source == "existing":
                rec_type = "Existing"
            elif source == "new":
                rec_type = "New"
            else:
                rec_type = "Rec."
            priority = (rec.get("priority") or "").replace("_", " ").title()
            rec_data.append([
                Paragraph(str(i), styles["CellText"]),
                Paragraph(rec_type, styles["CellText"]),
                Paragraph(priority, styles["CellText"]),
                Paragraph(rec.get("clause_text", ""), styles["CellText"]),
                Paragraph(rec.get("justification", ""), styles["CellText"]),
                Paragraph(rec.get("existing_coverage", "N/A"), styles["CellText"]),
            ])
        t = Table(rec_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(header_style)
        elements.append(t)
    else:
        elements.append(Paragraph("No recommendations generated.", styles["Body"]))

    # ── Footer endmark ──
    elements.append(Spacer(1, 10 * mm))
    elements.append(Paragraph(f"— End of Report — Generated on {now_str} —", styles["SmallGray"]))

    # ── Build PDF ──
    doc.build(elements, onFirstPage=_footer_header, onLaterPages=_footer_header)
    buf.seek(0)
    return buf
