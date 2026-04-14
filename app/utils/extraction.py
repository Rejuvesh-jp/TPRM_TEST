import io
import logging
from pathlib import Path

logger = logging.getLogger("tprm.extraction")


def extract_text_from_pdf(file_path: str | Path) -> str:
    """Extract text from a PDF file."""
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(pages)


def extract_text_from_pdf_bytes(content: bytes) -> str:
    """Extract text from PDF bytes."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(pages)


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract text from a DOCX file."""
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def extract_text_from_docx_bytes(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def extract_text_from_txt(file_path: str | Path) -> str:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_text_from_txt_bytes(content: bytes) -> str:
    """Extract text from plain text bytes."""
    return content.decode("utf-8", errors="replace")


def extract_text_from_xlsx(file_path: str | Path) -> str:
    """Extract text from an Excel file."""
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path), read_only=True, data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets)


def extract_text_from_xlsx_bytes(content: bytes) -> str:
    """Extract text from Excel bytes."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets)


def extract_text_from_csv(file_path: str | Path) -> str:
    """Extract text from a CSV file."""
    import csv

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = [" | ".join(row) for row in reader if any(row)]
    return "\n".join(rows)


def extract_text_from_json(file_path: str | Path) -> str:
    """Extract text from a JSON file — return pretty-printed content."""
    import json

    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, indent=2)


# Dispatcher
EXTRACTORS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".txt": extract_text_from_txt,
    ".csv": extract_text_from_csv,
    ".xlsx": extract_text_from_xlsx,
    ".json": extract_text_from_json,
}

BYTE_EXTRACTORS = {
    ".pdf": extract_text_from_pdf_bytes,
    ".docx": extract_text_from_docx_bytes,
    ".txt": extract_text_from_txt_bytes,
    ".xlsx": extract_text_from_xlsx_bytes,
}


def extract_text(file_path: str | Path) -> str:
    """Extract text from a file based on its extension."""
    ext = Path(file_path).suffix.lower()
    extractor = EXTRACTORS.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file type: {ext}")
    logger.info(f"Extracting text from {file_path} (type: {ext})")
    return extractor(file_path)


def extract_text_from_bytes(content: bytes, extension: str) -> str:
    """Extract text from file bytes based on extension."""
    ext = extension.lower()
    extractor = BYTE_EXTRACTORS.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file type for byte extraction: {ext}")
    return extractor(content)
